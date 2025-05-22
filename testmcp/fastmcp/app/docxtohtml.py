import docx
import html
import os
import io
import logging
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.numbering import CT_NumPr
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.shared import RGBColor, Pt
from docx.table import Table
from lxml import etree
from PIL import Image
from datetime import datetime
from docx.enum.text import WD_LINE_SPACING

# 로깅 설정
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('word_to_html_converter.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

def convert_rgb_to_hex(rgb_color):
    """RGB 색상을 HEX 코드로 변환"""
    if rgb_color is None:
        return None
    
    # RGBColor 객체는 .rgb 속성을 가지고 있으며, 이는 3바이트 정수입니다.
    # 각 색상 컴포넌트를 추출해야 합니다.
    if isinstance(rgb_color, RGBColor):
        # RGBColor 객체에서 RGB 값 추출 (6바이트 정수)
        rgb_int = rgb_color.rgb
        # RGB 값에서 각각의 컴포넌트 추출
        r = rgb_int & 0xFF
        g = (rgb_int >> 8) & 0xFF
        b = (rgb_int >> 16) & 0xFF
        return f'#{r:02x}{g:02x}{b:02x}'
    
    # 다른 형태의 RGB 컬러 객체 처리
    if hasattr(rgb_color, 'r') and hasattr(rgb_color, 'g') and hasattr(rgb_color, 'b'):
        return f'#{rgb_color.r:02x}{rgb_color.g:02x}{rgb_color.b:02x}'
        
    return None


def get_text_indentation(paragraph):
    """문단 들여쓰기 정보 추출"""
    paragraph_format = paragraph.paragraph_format

    style = {}

    first_line = paragraph_format.first_line_indent
    left_indent = paragraph_format.left_indent
    right_indent = paragraph_format.right_indent

    if first_line is not None and first_line.pt is not None:
        first_line_pt = first_line.pt
        style['text-indent'] = f'{first_line_pt}pt'
        if first_line_pt < 0:
            style['padding-left'] = f'{-first_line_pt}pt'

    if left_indent is not None and left_indent.pt is not None:
        left_pt = left_indent.pt
        existing = float(style.get('padding-left', '0').replace('pt', ''))
        style['padding-left'] = f'{existing + left_pt}pt'

    if right_indent is not None and right_indent.pt is not None:
        style['padding-right'] = f'{right_indent.pt}pt'

    return style


def get_line_spacing(paragraph):
    """줄 간격 정보 추출"""
    style = {}
    paragraph_format = paragraph.paragraph_format
    spacing = paragraph_format.line_spacing
    rule = paragraph_format.line_spacing_rule

    if spacing is not None:
        if rule == WD_LINE_SPACING.MULTIPLE:
            style["line-height"] = f"{spacing:.2f}"
        elif rule in (WD_LINE_SPACING.EXACTLY, WD_LINE_SPACING.AT_LEAST):
            style["line-height"] = f"{spacing:.2f}pt"
    
    return style


def get_paragraph_spacing(paragraph):
    """문단 간격 정보 추출"""
    style = {}
    paragraph_format = paragraph.paragraph_format
    
    before = paragraph_format.space_before
    after = paragraph_format.space_after
    
    if before and before.pt > 0:
        style['margin-top'] = f'{before.pt}pt'
    
    if after and after.pt > 0:
        style['margin-bottom'] = f'{after.pt}pt'
    
    return style


def get_list_info(paragraph, doc):
    """단락의 목록(번호 매기기 또는 글머리 기호) 정보를 추출"""
    try:
        if not hasattr(paragraph._element.pPr, 'numPr'):
            return {'is_list': False, 'list_type': None}
        
        num_pr = paragraph._element.pPr.numPr
        if num_pr is None:
            return {'is_list': False, 'list_type': None}
        
        num_id = None
        ilvl = None
        
        for child in num_pr.iterchildren():
            if child.tag.endswith('numId'):
                num_id = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            elif child.tag.endswith('ilvl'):
                ilvl = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
        
        if num_id is None or ilvl is None:
            return {'is_list': False, 'list_type': None}
        
        return {
            'is_list': True,
            'list_type': 'ul',
            'bullet': '•',
            'num_id': num_id,
            'level': int(ilvl)
        }
        
    except Exception as e:
        logger.error(f"목록 정보 추출 중 오류 발생: {str(e)}")
        return {'is_list': False, 'list_type': None}


FONT_MAPPING = {
    "Malgun Gothic": "'맑은 고딕', 'Malgun Gothic', sans-serif",
    "Batang": "'바탕', 'Batang', serif",
    "Gulim": "'굴림', 'Gulim', sans-serif",
    "Dotum": "'돋움', 'Dotum', sans-serif",
    "Gungsuh": "'궁서', 'Gungsuh', serif",
    "Arial": "'Arial', sans-serif",
    "Times New Roman": "'Times New Roman', serif",
    "Calibri": "'Calibri', sans-serif",
    "Cambria": "'Cambria', serif",
    "Georgia": "'Georgia', serif",
    "Verdana": "'Verdana', sans-serif",
    "Tahoma": "'Tahoma', sans-serif",
    "Apple SD Gothic Neo": "'Apple SD Gothic Neo', 'Noto Sans KR', sans-serif",
    "Noto Sans CJK KR": "'Noto Sans KR', sans-serif",
    "Noto Serif CJK KR": "'Noto Serif KR', serif"
}

USED_FONTS = set()  # 사용된 폰트 수집용

def get_font_family_style(font_name):
    USED_FONTS.add(font_name)  # 사용된 폰트 기록
    if font_name in FONT_MAPPING:
        return f"font-family: {FONT_MAPPING[font_name]}"
    else:
        return f"font-family: '{font_name}', 'Noto Sans KR', sans-serif"


def get_formatted_text_as_html(paragraph):
    """문단 내 서식이 적용된 텍스트를 HTML 태그 형식으로 추출 (모든 서식 포함)"""
    result = ""
    
    for run in paragraph.runs:
        lines = run.text.splitlines()
        for idx, line in enumerate(lines):
            text = html.escape(line)
            if not text and idx == len(lines) - 1:
                continue
            
            style_attrs = []
            
            if run.font.name:
                style_attrs.append(get_font_family_style(run.font.name))
            
            if run.font.size:
                try:
                    size_pt = run.font.size.pt
                    style_attrs.append(f"font-size: {size_pt}pt")
                except AttributeError:
                    if isinstance(run.font.size, int) or isinstance(run.font.size, float):
                        style_attrs.append(f"font-size: {run.font.size/2}pt")
            
            if run.bold:
                style_attrs.append("font-weight: bold")
            elif hasattr(run.font, 'cs_bold') and run.font.cs_bold:
                style_attrs.append("font-weight: bold")
            
            if hasattr(run.font, 'color') and run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                if len(str(rgb)) == 6:
                    style_attrs.append(f"color: #{rgb}")
                elif len(str(rgb)) == 8:
                    style_attrs.append(f"color: #{rgb[2:]}")
            
            if run.font.highlight_color:
                highlight_colors = {
                    'YELLOW (7)': '#FFFF00',
                    'BRIGHT_GREEN (2)': '#00FF00',
                    'TURQUOISE (3)': '#00FFFF',
                    'PINK (5)': '#FF00FF',
                    'BLUE (1)': '#0000FF',
                    'RED (6)': '#FF0000',
                    'DARK_BLUE (4)': '#000080',
                    'TEAL (8)': '#008080',
                    'GREEN (9)': '#008000',
                    'VIOLET (10)': '#800080',
                    'DARK_RED (11)': '#800000',
                    'DARK_YELLOW (12)': '#808000',
                    'GRAY_25 (13)': '#E6E6E6',
                    'GRAY_50 (14)': '#808080',
                    'BLACK (15)': '#000000',
                    'WHITE (16)': '#FFFFFF'
                }
                highlight_key = run.font.highlight_color
                if isinstance(highlight_key, int):
                    index_to_name = {
                        1: 'BLUE (1)',
                        2: 'BRIGHT_GREEN (2)',
                        3: 'TURQUOISE (3)',
                        4: 'DARK_BLUE (4)',
                        5: 'PINK (5)',
                        6: 'RED (6)',
                        7: 'YELLOW (7)',
                        8: 'TEAL (8)',
                        9: 'GREEN (9)',
                        10: 'VIOLET (10)',
                        11: 'DARK_RED (11)',
                        12: 'DARK_YELLOW (12)',
                        13: 'GRAY_25 (13)',
                        14: 'GRAY_50 (14)',
                        15: 'BLACK (15)',
                        16: 'WHITE (16)'
                    }
                    if highlight_key in index_to_name:
                        highlight_key = index_to_name[highlight_key]
                
                if highlight_key in highlight_colors:
                    style_attrs.append(f"background-color: {highlight_colors[highlight_key]}")
            
            if run.italic or (hasattr(run.font, 'cs_italic') and run.font.cs_italic):
                style_attrs.append("font-style: italic")
            if run.underline:
                style_attrs.append("text-decoration: underline")
            if run.font.strike:
                style_attrs.append("text-decoration: line-through")
            if run.font.superscript:
                style_attrs.append("vertical-align: super")
                style_attrs.append("font-size: smaller")
            if run.font.subscript:
                style_attrs.append("vertical-align: sub")
                style_attrs.append("font-size: smaller")
            
            if hasattr(run.font, 'spacing') and run.font.spacing:
                spacing_pt = run.font.spacing / 20.0
                style_attrs.append(f"letter-spacing: {spacing_pt}pt")
            
            if style_attrs:
                style_str = "; ".join(style_attrs)
                text = f'<span style="{style_str}">{text}</span>'
            elif text:
                text = f'<span>{text}</span>'
            
            result += text
            if idx < len(lines) - 1:
                result += "<br>"
    
    return result


def get_paragraph_alignment(paragraph):
    """문단 정렬을 HTML 스타일로 변환"""
    alignment = paragraph.alignment
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return 'center'
    elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        return 'right'
    elif alignment == WD_ALIGN_PARAGRAPH.LEFT:
        return 'left'
    elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return 'justify'
    return 'left'  # 기본값


def get_merged_cells_info(table):
    """표의 병합된 셀 정보를 더 정확하게 가져옵니다"""
    merged_cells = {}
    
    # 모든 셀의 기본 정보 초기화
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            merged_cells[(i, j)] = {'rowspan': 1, 'colspan': 1, 'skip': False}
    
    # 가로 병합(colspan) 처리
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if cell._element.tcPr is not None:
                gridSpan = cell._element.tcPr.xpath('.//w:gridSpan')
                if gridSpan:
                    colspan = int(gridSpan[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'))
                    if colspan > 1:
                        merged_cells[(i, j)]['colspan'] = colspan
                        for k in range(1, colspan):
                            if (i, j + k) in merged_cells:
                                merged_cells[(i, j + k)]['skip'] = True
    
    # 세로 병합(rowspan) 처리
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if cell._element.tcPr is not None:
                vMerge = cell._element.tcPr.xpath('.//w:vMerge')
                if vMerge:
                    vMergeVal = vMerge[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    if vMergeVal == 'restart':
                        rowspan = 1
                        for k in range(i + 1, len(table.rows)):
                            next_cell = table.rows[k].cells[j]
                            if (next_cell._element.tcPr is not None and 
                                next_cell._element.tcPr.xpath('.//w:vMerge') and 
                                next_cell._element.tcPr.xpath('.//w:vMerge')[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') is None):
                                rowspan += 1
                                merged_cells[(k, j)]['skip'] = True
                            else:
                                break
                        merged_cells[(i, j)]['rowspan'] = rowspan
    
    return merged_cells


def convert_border_style(val, sz=None):
    """Word 테두리 스타일을 CSS 스타일로 변환"""
    if val is None or val == 'nil' or val == 'none':
        return 'none'
    
    # 기본 스타일 매핑 - 지원 안하는 것 유사하게 매핑
    style_map = {
        'single': lambda size: 'solid',
        'thick': lambda size: 'solid',
        'double': lambda size: 'double',
        'dotted': lambda size: 'dotted',
        'dashed': lambda size: 'dashed',
        'dotDash': lambda size: size and float(size) > 4 and 'dashed' or 'dotted',
        'dotDotDash': lambda size: 'dashed',
        'triple': lambda size: 'double',  # CSS에서 triple 지원 안함
        'thinThickSmallGap': lambda size: 'double',
        'thickThinSmallGap': lambda size: 'double',
        'thinThickThinSmallGap': lambda size: 'double',
        'thinThickMediumGap': lambda size: 'double',
        'thickThinMediumGap': lambda size: 'double',
        'thinThickThinMediumGap': lambda size: 'double',
        'thinThickLargeGap': lambda size: 'double',
        'thickThinLargeGap': lambda size: 'double',
        'thinThickThinLargeGap': lambda size: 'double',
        'wave': lambda size: 'solid',
        'doubleWave': lambda size: 'double',
        'dashSmallGap': lambda size: 'dashed',
        'dashDotStroked': lambda size: 'dashed',
        'threeDEmboss': lambda size: 'ridge',
        'threeDEngrave': lambda size: 'groove',
        'outset': lambda size: 'outset',
        'inset': lambda size: 'inset'
    }
    
    style_func = style_map.get(val, lambda size: 'solid')
    return style_func(sz)


def get_border_width(sz, style=None):
    """테두리 두께를 포인트로 변환"""
    if sz is None:
        return '0'
    
    # Word의 테두리 두께는 1/8 포인트 단위
    width = float(sz) / 8
    
    # 특정 스타일에 따른 두께 조정
    if style in ['double', 'triple']:
        width = max(width, 2.25)  # 최소 2.25pt
    elif style in ['dotted', 'dashed']:
        width = min(width, 1.5)   # 최대 1.5pt
    
    return f'{width:.1f}pt'


def get_border_color(color, val=None):
    """테두리 색상을 CSS 색상으로 변환"""
    if color is None or color == 'auto':
        if val in ['threeDEngrave', 'threeDEmboss']:
            return '#808080'  # 3D 효과를 위한 회색
        return '#000000'
    return f'#{color}'


def get_table_cell_styles(cell, row_index, col_index, total_rows, total_cols):
    """테이블 셀의 스타일 정보 추출"""
    style_attrs = []
    
    # 셀 배경색 추출
    if cell._element.tcPr is not None:
        shading = cell._element.tcPr.xpath('./w:shd')
        if shading:
            fill_color = shading[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
            if fill_color and fill_color != 'auto':
                style_attrs.append(f'background-color: #{fill_color}')
            
            # 음영 효과 처리
            shading_pattern = shading[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            if shading_pattern and shading_pattern != 'clear':
                style_attrs.append('background-image: linear-gradient(45deg, rgba(0,0,0,0.1) 25%, transparent 25%, transparent 50%, rgba(0,0,0,0.1) 50%, rgba(0,0,0,0.1) 75%, transparent 75%, transparent)')
                style_attrs.append('background-size: 10px 10px')
    
    # 테두리 스타일 추출
    if cell._element.tcPr is not None:
        borders = {}
        border_elements = cell._element.tcPr.xpath('./w:tcBorders/*')
        
        # 기본값 설정 - 테이블의 내부 테두리 스타일을 사용
        borders = {
            'top': {'use_inside': row_index > 0},
            'right': {'use_inside': col_index < total_cols - 1},
            'bottom': {'use_inside': row_index < total_rows - 1},
            'left': {'use_inside': col_index > 0}
        }
        
        # 실제 테두리 정보 추출
        for border in border_elements:
            side = border.tag.split('}')[-1]
            val = border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            sz = border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz')
            
            if val == 'nil' or val == 'none':
                borders[side].update({'style': 'none', 'width': '0', 'color': 'transparent', 'space': '0'})
            else:
                style = convert_border_style(val, sz)
                width = get_border_width(sz, val)
                color = get_border_color(border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color'), val)
                space = border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', '0')
                
                borders[side].update({
                    'style': style,
                    'width': width,
                    'color': color,
                    'space': space
                })
        
        # 각 방향별 테두리 스타일 적용
        for side, border in borders.items():
            if border.get('style') == 'none' and border.get('use_inside', False):
                # 내부 테두리 스타일 사용
                if side in ['top', 'bottom']:
                    style_attrs.append(f'border-{side}: var(--table-border-insideH)')
                else:  # left, right
                    style_attrs.append(f'border-{side}: var(--table-border-insideV)')
            elif border.get('style', 'none') != 'none':
                if border.get('style') in ['ridge', 'groove', 'inset', 'outset']:
                    style_attrs.append(f'border-{side}-style: {border["style"]}')
                    style_attrs.append(f'border-{side}-width: {border["width"]}')
                    style_attrs.append(f'border-{side}-color: {border["color"]}')
                else:
                    style_attrs.append(f'border-{side}: {border["width"]} {border["style"]} {border["color"]}')
                
                # 테두리 간격이 있는 경우
                if float(border.get('space', '0')) > 0:
                    space_pt = float(border['space']) / 20
                    style_attrs.append(f'padding-{side}: {space_pt}pt')
            else:
                style_attrs.append(f'border-{side}: none')

    # 셀 너비 추출
    if cell._element.tcPr is not None:
        width = cell._element.tcPr.xpath('./w:tcW')
        if width:
            w_value = width[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
            w_type = width[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'dxa')
            if w_type == 'dxa':  # twips 단위 (1/20pt)
                cell_width = float(w_value) / 20  # 트위프를 포인트로 변환
                style_attrs.append(f'width: {cell_width}pt')
            elif w_type == 'pct':  # 퍼센트 (1/50%) 단위)
                cell_width = float(w_value) / 50
                style_attrs.append(f'width: {cell_width}%')
    
    # 수직 정렬 추출
    if cell._element.tcPr is not None:
        v_align = cell._element.tcPr.xpath('./w:vAlign')
        if v_align:
            align_value = v_align[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            if align_value == 'center':
                style_attrs.append('vertical-align: middle')
            elif align_value == 'bottom':
                style_attrs.append('vertical-align: bottom')
            else:
                style_attrs.append('vertical-align: top')
    
    # 기본 패딩 설정 (테두리 간격이 없는 경우)
    if not any('padding-' in attr for attr in style_attrs):
        style_attrs.append('padding: 2pt')
    
    # 텍스트 줄바꿈 및 오버플로우 처리
    style_attrs.append('word-wrap: break-word')
    style_attrs.append('white-space: normal')
    style_attrs.append('overflow-wrap: break-word')
    style_attrs.append('hyphens: auto')
    
    return "; ".join(style_attrs)


def get_table_styles(table):
    """테이블 전체 스타일 정보 추출"""
    style_attrs = []
    
    if table._element.tblPr is not None:
        # 테이블 너비 추출
        width = table._element.tblPr.xpath('./w:tblW')
        if width:
            w_value = width[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
            w_type = width[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'dxa')
            if w_type == 'dxa':  # twips 단위
                table_width = float(w_value) / 20  # twips를 포인트로 변환
                style_attrs.append(f'width: {table_width}pt')
            elif w_type == 'pct':  # 퍼센트 단위
                table_width = float(w_value) / 50
                style_attrs.append(f'width: {table_width}%')
            elif w_type == 'auto':
                style_attrs.append('width: auto')
        else:
            # 너비가 지정되지 않은 경우 자동으로 설정
            style_attrs.append('width: auto')
        
        # 테이블 여백 추출 -> 셀 간 간격
        margins = table._element.tblPr.xpath('./w:tblCellSpacing')
        if margins:
            margin_value = float(margins[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', '0')) / 20
            style_attrs.append(f'border-spacing: {margin_value}pt')
            style_attrs.append('border-collapse: separate')
        else:
            style_attrs.append('border-collapse: collapse')
        
        # 테이블 정렬 추출
        jc = table._element.tblPr.xpath('./w:jc')
        if jc:
            align_value = jc[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            if align_value == 'center':
                style_attrs.append('margin-left: auto')
                style_attrs.append('margin-right: auto')
            elif align_value == 'right':
                style_attrs.append('margin-left: auto')
                style_attrs.append('margin-right: 0')
        
        # 테이블 전체 테두리 스타일 추출
        borders = {}
        border_elements = table._element.tblPr.xpath('.//w:tblBorders/*')
        
        # 테이블 전체 테두리 기본값 설정
        for side in ['top', 'right', 'bottom', 'left', 'insideH', 'insideV']:
            borders[side] = {'style': 'none', 'width': '0', 'color': 'transparent'}
        
        # 실제 테두리 정보 추출
        for border in border_elements:
            side = border.tag.split('}')[-1]
            val = border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            sz = border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz')
            
            if val == 'nil' or val == 'none':
                borders[side] = {'style': 'none', 'width': '0', 'color': 'transparent'}
            else:
                style = convert_border_style(val, sz)
                width = get_border_width(sz, val)
                color = get_border_color(border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color'), val)
                borders[side] = {'style': style, 'width': width, 'color': color}
        
        # 테이블 전체 테두리 스타일 적용
        for side in ['top', 'right', 'bottom', 'left']:
            border = borders.get(side)
            if border['style'] != 'none':
                style_attrs.append(f'border-{side}: {border["width"]} {border["style"]} {border["color"]}')
            else:
                style_attrs.append(f'border-{side}: none')
        
        # 내부 테두리 스타일을 CSS 변수로 저장 (셀에서 사용) -> 이거 중복 같은데
        for side in ['insideH', 'insideV']:
            border = borders.get(side)
            if border['style'] != 'none':
                style_attrs.append(f'--table-border-{side}: {border["width"]} {border["style"]} {border["color"]}')
            else:
                style_attrs.append(f'--table-border-{side}: none')
    
    return "; ".join(style_attrs)


def detect_toc(paragraph):
    """목차 여부 감지"""
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    if "toc" in style_name or "목차" in style_name:
        return True
    
    for run in paragraph.runs:
        if hasattr(run, '_element') and run._element.xpath('.//w:fldChar'):
            return True
    
    return False


def get_header_footer_content(section):
    """섹션의 머리글과 바닥글 내용을 추출"""
    content = {
        'header': [],
        'footer': []
    }
    
    # 머리글 처리
    if section.header is not None:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                style = []
                alignment = get_paragraph_alignment(paragraph)
                indentation = get_text_indentation(paragraph)
                spacing = get_paragraph_spacing(paragraph)
                
                style.append(f"text-align: {alignment}")
                for key, value in indentation.items():
                    style.append(f"{key}: {value}")
                for key, value in spacing.items():
                    style.append(f"{key}: {value}")
                
                style_str = "; ".join(style)
                content['header'].append(
                    f'<p style="{style_str}">{get_formatted_text_as_html(paragraph)}</p>'
                )
        
        # 머리글의 표 처리
        for table in section.header.tables:
            table_html = process_table(table)
            content['header'].append(table_html)
    
    # 바닥글 처리
    if section.footer is not None:
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                style = []
                alignment = get_paragraph_alignment(paragraph)
                indentation = get_text_indentation(paragraph)
                spacing = get_paragraph_spacing(paragraph)
                
                style.append(f"text-align: {alignment}")
                for key, value in indentation.items():
                    style.append(f"{key}: {value}")
                for key, value in spacing.items():
                    style.append(f"{key}: {value}")
                
                style_str = "; ".join(style)
                content['footer'].append(
                    f'<p style="{style_str}">{get_formatted_text_as_html(paragraph)}</p>'
                )
        
        # 바닥글의 표 처리
        for table in section.footer.tables:
            table_html = process_table(table)
            content['footer'].append(table_html)
    
    return content

def process_paragraph(paragraph, doc):
    """단락 처리"""
    try:
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        alignment = get_paragraph_alignment(paragraph)
        indentation = get_text_indentation(paragraph)
        spacing = get_paragraph_spacing(paragraph)
        line_spacing = get_line_spacing(paragraph)
        
        # 스타일 속성 목록
        style_attrs = [f"text-align: {alignment}"]
        
        # 들여쓰기 추가
        for key, value in indentation.items():
            style_attrs.append(f"{key}: {value}")
        
        # 여백 추가
        for key, value in spacing.items():
            style_attrs.append(f"{key}: {value}")
        
        # 줄 간격 추가
        if line_spacing:
            style_attrs.append(f"line-height: {line_spacing}")
        
        # 스타일 문자열 생성
        style_str = "; ".join(style_attrs)
        
        # 페이지 나누기 확인
        has_page_break = False
        if hasattr(paragraph._element, 'lastRenderedPageBreak'):
            has_page_break = True
        
        # 명시적 페이지 나누기 확인
        if paragraph._p.xpath('.//w:br[@w:type="page"]'):
            has_page_break = True
        
        # 페이지 나누기가 있는 경우 처리
        if has_page_break:
            return f'<div class="page-break"></div>\n<p style="{style_str}">{get_formatted_text_as_html(paragraph)}</p>\n'
        
        # 목차 감지 및 처리
        is_toc = detect_toc(paragraph)
        if is_toc:
            return f'<div class="toc" style="{style_str}">{get_formatted_text_as_html(paragraph)}</div>\n'
        
        # 빈 단락 처리
        text = paragraph.text.strip() if paragraph.text else ""
        if not text:
            return f'<p style="{style_str}"><br></p>\n'
        
        # 제목 스타일 처리
        if "heading" in style_name:
            if "1" in style_name:
                level = 1
            elif "2" in style_name:
                level = 2
            elif "3" in style_name:
                level = 3
            else:
                level = 4
            
            # 안전한 ID 생성
            safe_text = ''.join(c for c in text.lower() if c.isalnum() or c in '-_ ')[:20]
            heading_id = f"heading-{level}-{safe_text}"
            return f'<h{level} id="{heading_id}" style="{style_str}">{get_formatted_text_as_html(paragraph)}</h{level}>\n'
        
        # 일반 단락 처리
        return f'<p style="{style_str}">{get_formatted_text_as_html(paragraph)}</p>\n'
    
    except Exception as e:
        logger.error(f"단락 처리 중 오류 발생: {str(e)}")
        return ""

def process_comment(comment_element):
    """주석 처리"""
    try:
        author = comment_element.get('w:author', 'Unknown')
        date = comment_element.get('w:date', '')
        comment_text = ""
        
        # 주석 내용 추출
        for para in comment_element.xpath('.//w:p'):
            comment_text += get_formatted_text_as_html(para)
        
        return f'''
            <span class="comment" title="Comment by {html.escape(author)} on {date}">
                <span class="comment-marker">💬</span>
                <span class="comment-text">{comment_text}</span>
            </span>
        '''
    except Exception as e:
        logger.error(f"주석 처리 중 오류 발생: {str(e)}")
        return ""

def process_revision(revision_element):
    """변경 내역 처리"""
    try:
        author = revision_element.get('w:author', 'Unknown')
        date = revision_element.get('w:date', '')
        revision_type = revision_element.tag.split('}')[-1]
        
        if revision_type == 'ins':
            return f'<ins class="revision" title="Added by {author} on {date}">'
        elif revision_type == 'del':
            return f'<del class="revision" title="Deleted by {author} on {date}">'
        
        return ""
    except Exception as e:
        logger.error(f"변경 내역 처리 중 오류 발생: {str(e)}")
        return ""

def process_image(doc, image_element, image_counter, images_dir):
    """이미지 처리 및 저장"""
    try:
        # 이미지 관련 요소 찾기
        blip = image_element.xpath('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
        if not blip:
            return ""
        
        image_rid = blip[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if not image_rid:
            return ""
        
        # 이미지 데이터 가져오기
        image_part = doc.part.related_parts[image_rid]
        image_bytes = image_part.blob
        
        # 이미지 포맷 확인
        image = Image.open(io.BytesIO(image_bytes))
        format_ext = image.format.lower()
        
        # 이미지 파일명 생성
        image_filename = f"image_{image_counter}.{format_ext}"
        image_path = os.path.join(images_dir, image_filename)
        
        # 이미지 저장
        with open(image_path, "wb") as img_file:
            img_file.write(image_bytes)
        
        # 이미지 크기 정보 추출
        extent = image_element.xpath('.//wp:extent', namespaces={'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
        style = ""
        if extent:
            cx = int(extent[0].get('cx', 0)) / 9525  # EMU to points
            cy = int(extent[0].get('cy', 0)) / 9525
            style = f"width: {cx}pt; height: {cy}pt;"
        
        # HTML img 태그 생성
        return f'<img src="{image_filename}" alt="Document image {image_counter}" style="{style}" />'
    
    except Exception as e:
        logger.error(f"이미지 처리 중 오류 발생: {str(e)}")
        return ""

def process_equation(equation_element):
    """수식 처리"""
    try:
        mathml = str(equation_element)
        return f'<math xmlns="http://www.w3.org/1998/Math/MathML">{mathml}</math>'
    except Exception as e:
        logger.error(f"수식 처리 중 오류 발생: {str(e)}")
        return ""

def get_table_properties(table):
    """표의 속성 정보를 추출"""
    properties = {}
    
    if table._element.tblPr is not None:
        # 표 정렬
        jc = table._element.tblPr.xpath('./w:jc')
        if jc:
            align_value = jc[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            properties['alignment'] = align_value  # left, center, right
        
        # 표 들여쓰기
        ind = table._element.tblPr.xpath('./w:tblInd')
        if ind:
            ind_type = ind[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
            ind_value = ind[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
            if ind_type == 'dxa':  # 트위프 단위
                properties['indent'] = f'{float(ind_value)/20}pt'  # 트위프를 포인트로 변환
            elif ind_type == 'pct':  # 퍼센트 단위
                properties['indent'] = f'{float(ind_value)/50}%'
            elif ind_type == 'auto':
                properties['indent'] = '0'
        
        # 표 너비
        width = table._element.tblPr.xpath('./w:tblW')
        if width:
            w_type = width[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
            w_value = width[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
            if w_type == 'dxa':
                properties['width'] = f'{float(w_value)/20}pt'
            elif w_type == 'pct':
                properties['width'] = f'{float(w_value)/50}%'
            elif w_type == 'auto':
                properties['width'] = 'auto'
        
        # 표 여백
        margins = table._element.tblPr.xpath('./w:tblCellSpacing')
        if margins:
            margin_value = float(margins[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', '0')) / 20
            properties['cell-spacing'] = f'{margin_value}pt'
        
        # 표 레이아웃
        layout = table._element.tblPr.xpath('./w:tblLayout')
        if layout:
            layout_type = layout[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
            properties['layout'] = layout_type  # fixed, autofit
        
        # 표 텍스트 배치
        text_flow = table._element.tblPr.xpath('./w:tblpPr')
        if text_flow:
            properties['text-flow'] = 'around'  # 둘러싸기
        else:
            properties['text-flow'] = 'none'  # 없음
        
        # 표 머리글 행 반복
        header = table._element.tblPr.xpath('./w:tblHeader')
        if header:
            properties['header-rows'] = 'repeat'
        
        # 표 페이지 나누기 설정
        cant_split = table._element.tblPr.xpath('./w:tblpPr/w:cantSplit')
        if cant_split:
            properties['page-break-inside'] = 'avoid'
    
    return properties

def process_table(table):
    """테이블을 HTML로 변환"""
    try:
        # 테이블 스타일 추출
        table_style = get_table_styles(table)
        
        # 테이블 속성 추출
        table_properties = get_table_properties(table)
        
        # 스타일 속성 목록 생성
        style_attrs = []
        
        # 정렬
        if 'alignment' in table_properties:
            if table_properties['alignment'] == 'center':
                style_attrs.append('margin-left: auto; margin-right: auto')
            elif table_properties['alignment'] == 'right':
                style_attrs.append('margin-left: auto; margin-right: 0')
        
        # 들여쓰기
        if 'indent' in table_properties:
            if table_properties['indent'] != '0':
                style_attrs.append(f'margin-left: {table_properties["indent"]}')
        
        # 너비
        if 'width' in table_properties:
            style_attrs.append(f'width: {table_properties["width"]}')
        
        # 셀 간격
        if 'cell-spacing' in table_properties:
            style_attrs.append(f'border-spacing: {table_properties["cell-spacing"]}')
            style_attrs.append('border-collapse: separate')
        else:
            style_attrs.append('border-collapse: collapse')
        
        # 레이아웃
        if 'layout' in table_properties:
            style_attrs.append(f'table-layout: {table_properties["layout"]}')
        
        # 텍스트 배치
        if 'text-flow' in table_properties:
            if table_properties['text-flow'] == 'around':
                style_attrs.append('float: left')
                style_attrs.append('margin-right: 1cm')
        
        # 페이지 나누기
        if 'page-break-inside' in table_properties:
            style_attrs.append('page-break-inside: avoid')
        
        # 기존 테이블 스타일과 결합
        if style_attrs:
            table_style = f"{table_style}; {'; '.join(style_attrs)}"
        
        # 병합된 셀 정보 가져오기
        merged_cells = get_merged_cells_info(table)
        
        # 테이블 HTML 시작
        html = f'<table style="{table_style}">\n'
        
        # 행 처리
        for i, row in enumerate(table.rows):
            html += '<tr>\n'
            
            # 열 처리
            for j, cell in enumerate(row.cells):
                # 병합된 셀인 경우 건너뛰기
                if merged_cells.get((i, j), {}).get('skip', False):
                    continue
                
                # 셀 스타일 추출
                cell_style = get_table_cell_styles(cell, i, j, len(table.rows), len(row.cells))
                
                # 병합 정보 추가
                rowspan = merged_cells.get((i, j), {}).get('rowspan', 1)
                colspan = merged_cells.get((i, j), {}).get('colspan', 1)
                
                # 셀 HTML 생성
                html += f'<td style="{cell_style}"'
                if rowspan > 1:
                    html += f' rowspan="{rowspan}"'
                if colspan > 1:
                    html += f' colspan="{colspan}"'
                html += '>'
                
                # 셀 내용 처리
                for paragraph in cell.paragraphs:
                    html += process_paragraph(paragraph, table._parent)
                
                html += '</td>\n'
            
            html += '</tr>\n'
        
        html += '</table>\n'
        return html
        
    except Exception as e:
        logger.error(f"테이블 처리 중 오류 발생: {str(e)}")
        return ""

def generate_css():
    """스타일시트 생성"""
    return """
    @charset "UTF-8";
    
    /* 기본 스타일 */
    :root {
        --main-font: 'Calibri', 'Malgun Gothic', sans-serif;
        --main-color: #333;
        --page-width: 21cm;
        --page-height: 29.7cm;
        --margin-top: 2.54cm;
        --margin-bottom: 2.54cm;
        --margin-left: 2.54cm;
        --margin-right: 2.54cm;
        --content-width: calc(var(--page-width) - var(--margin-left) - var(--margin-right));
    }
    
    body { 
        font-family: var(--main-font);
        font-size: 11pt;
        line-height: 1.15;
        color: var(--main-color);
        margin: 0;
        padding: 0;
        background-color: #f0f0f0;
    }
    
    /* 문서 컨테이너 */
    .word-document { 
        width: var(--page-width);
        min-height: var(--page-height);
        padding: var(--margin-top) var(--margin-right) var(--margin-bottom) var(--margin-left);
        margin: 1cm auto;
        background: white;
        box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
        position: relative;
        box-sizing: border-box;
    }
    
    /* 단락 스타일 */
    p {
        margin: 0;
        padding: 0;
        word-wrap: break-word;
        margin-bottom: 8pt;
    }
    
    /* 제목 스타일 */
    h1 { font-size: 16pt; margin: 24pt 0 12pt 0; }
    h2 { font-size: 14pt; margin: 20pt 0 10pt 0; }
    h3 { font-size: 12pt; margin: 16pt 0 8pt 0; }
    h4 { font-size: 11pt; margin: 14pt 0 7pt 0; }
    
    /* 표 스타일 */
    table { 
        margin: 12pt 0;
        border-collapse: collapse;
        table-layout: fixed;
    }
    
    td { 
        padding: 6pt;
        border: 1px solid #ddd;
        word-wrap: normal;
        white-space: nowrap;
        overflow: visible;
        vertical-align: top;
        box-sizing: border-box;
    }
    
    /* 이미지 */
    img { 
        max-width: 100%;
        height: auto;
        margin: 8pt 0;
        display: block;
    }
    
    /* 인쇄용 스타일 */
    @media print {
        @page {
            size: {orientation};
            margin: 0;
        }
        
        body {
            margin: 0;
            padding: 0;
            background: none;
        }
        
        .word-document {
            width: 100%;
            min-height: 100%;
            padding: var(--margin-top) var(--margin-right) var(--margin-bottom) var(--margin-left);
            margin: 0;
            box-shadow: none;
            box-sizing: border-box;
            print-color-adjust: exact;
            -webkit-print-color-adjust: exact;
        }
        
        /* 페이지 나누기 */
        .page-break {
            page-break-before: always;
        }
    }
    """

def process_header_footer_table(table, doc):
    """머리글/바닥글의 표를 HTML로 변환 (병합된 셀 처리 개선)"""
    try:
        # 테이블 스타일 추출
        table_style = get_table_styles(table)
        table_style += "; border-collapse: collapse; width: 100%"
        
        # 병합된 셀 정보 가져오기 (본문 표 처리와 동일한 방식 사용)
        merged_cells = get_merged_cells_info(table)
        
        # 테이블 HTML 시작
        html = f'<table style="{table_style}">\n'
        
        # 행 처리
        for i, row in enumerate(table.rows):
            html += '<tr>\n'
            
            # 열 처리
            for j, cell in enumerate(row.cells):
                # 병합된 셀인 경우 건너뛰기
                if merged_cells.get((i, j), {}).get('skip', False):
                    continue
                
                # 셀 스타일 추출
                cell_style = get_table_cell_styles(cell, i, j, len(table.rows), len(row.cells))
                cell_style += "; border: 1px solid black; padding: 5px"
                
                # 병합 정보 추가
                rowspan = merged_cells.get((i, j), {}).get('rowspan', 1)
                colspan = merged_cells.get((i, j), {}).get('colspan', 1)
                
                # 셀 HTML 생성
                html += f'<td style="{cell_style}"'
                if rowspan > 1:
                    html += f' rowspan="{rowspan}"'
                if colspan > 1:
                    html += f' colspan="{colspan}"'
                html += '>'
                
                # 셀 내용 처리
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        style = []
                        alignment = get_paragraph_alignment(paragraph)
                        indentation = get_text_indentation(paragraph)
                        spacing = get_paragraph_spacing(paragraph)
                        
                        style.append(f"text-align: {alignment}")
                        for key, value in indentation.items():
                            style.append(f"{key}: {value}")
                        for key, value in spacing.items():
                            style.append(f"{key}: {value}")
                        
                        style_str = "; ".join(style)
                        html += f'<p style="{style_str}">{get_formatted_text_as_html(paragraph)}</p>'
                    else:
                        # 빈 단락의 경우 공백 추가
                        html += '<p>&nbsp;</p>'
                
                html += '</td>\n'
            
            html += '</tr>\n'
        
        html += '</table>\n'
        return html
        
    except Exception as e:
        logger.error(f"머리글/바닥글 표 처리 중 오류 발생: {str(e)}")
        return ""

def get_unique_header_footer(doc):
    """문서의 고유한 머리글/바닥글 내용을 추출"""
    unique_headers = {}
    unique_footers = {}
    
    for section in doc.sections:
        if section.header is not None:
            header_content = []
            header_key = ""
            
            for paragraph in section.header.paragraphs:
                header_key += paragraph.text
            
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        header_key += "".join(paragraph.text for paragraph in cell.paragraphs)
            
            if header_key not in unique_headers:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        style = []
                        alignment = get_paragraph_alignment(paragraph)
                        indentation = get_text_indentation(paragraph)
                        spacing = get_paragraph_spacing(paragraph)
                        
                        style.append(f"text-align: {alignment}")
                        for key, value in indentation.items():
                            style.append(f"{key}: {value}")
                        for key, value in spacing.items():
                            style.append(f"{key}: {value}")
                        
                        style_str = "; ".join(style)
                        header_content.append(
                            f'<p style="{style_str}">{get_formatted_text_as_html(paragraph)}</p>'
                        )
                
                for table in section.header.tables:
                    header_content.append(process_header_footer_table(table, doc))
                
                unique_headers[header_key] = header_content
        
        if section.footer is not None:
            footer_content = []
            footer_key = ""
            
            for paragraph in section.footer.paragraphs:
                footer_key += paragraph.text
            
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        footer_key += "".join(paragraph.text for paragraph in cell.paragraphs)
            
            if footer_key not in unique_footers:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        style = []
                        alignment = get_paragraph_alignment(paragraph)
                        indentation = get_text_indentation(paragraph)
                        spacing = get_paragraph_spacing(paragraph)
                        
                        style.append(f"text-align: {alignment}")
                        for key, value in indentation.items():
                            style.append(f"{key}: {value}")
                        for key, value in spacing.items():
                            style.append(f"{key}: {value}")
                        
                        style_str = "; ".join(style)
                        footer_content.append(
                            f'<p style="{style_str}">{get_formatted_text_as_html(paragraph)}</p>'
                        )
                
                for table in section.footer.tables:
                    footer_content.append(process_header_footer_table(table, doc))
                
                unique_footers[footer_key] = footer_content
    
    return list(unique_headers.values()), list(unique_footers.values())

def read_docx_as_html_structure(docx_path):
    """Word 문서를 HTML 구조로 변환"""
    try:
        logger.info(f"문서 변환 시작: {docx_path}")
        
        # 이미지 디렉토리 생성
        images_dir = "images"
        if not os.path.exists(images_dir):
            os.makedirs(images_dir)
        
        # docx 문서 열기
        doc = docx.Document(docx_path)
        
        # 머리글과 바닥글을 상단에 한 번 추출
        header_footer_html = ""
        image_counter = 0
        
        # 고유한 머리글과 바닥글 추출
        # unique_headers, unique_footers = get_unique_header_footer(doc)
        
        # 머리글과 바닥글 HTML 생성
        # if unique_headers or unique_footers:
        #     header_footer_html += '<div class="header-footer-preview">\n'
        #     
        #     if unique_headers:
        #         header_footer_html += '<div class="headers-preview">\n'
        #         header_footer_html += '<h3>머리글 미리보기</h3>\n'
        #         for i, header in enumerate(unique_headers):
        #             header_footer_html += f'<div class="header-section">\n'
        #             if len(unique_headers) > 1:
        #                 header_footer_html += f'<h4>머리글 {i+1}</h4>\n'
        #             header_footer_html += '\n'.join(header)
        #             header_footer_html += '\n</div>\n'
        #         header_footer_html += '</div>\n'
        #     
        #     if unique_footers:
        #         header_footer_html += '<div class="footers-preview">\n'
        #         header_footer_html += '<h3>바닥글 미리보기</h3>\n'
        #         for i, footer in enumerate(unique_footers):
        #             header_footer_html += f'<div class="footer-section">\n'
        #             if len(unique_footers) > 1:
        #                 header_footer_html += f'<h4>바닥글 {i+1}</h4>\n'
        #             header_footer_html += '\n'.join(footer)
        #             header_footer_html += '\n</div>\n'
        #         header_footer_html += '</div>\n'
        #     
        #     header_footer_html += '</div>\n'
        
        # 메타데이터 추출
        core_properties = doc.core_properties
        metadata = {
            "title": core_properties.title if core_properties.title else "Untitled",
            "author": core_properties.author if core_properties.author else "Unknown",
            "created": core_properties.created if core_properties.created else datetime.now(),
            "modified": core_properties.modified if core_properties.modified else datetime.now(),
            "last_modified_by": core_properties.last_modified_by if core_properties.last_modified_by else "Unknown",
            "revision": core_properties.revision if core_properties.revision else 1
        }

        # 첫 번째 섹션의 레이아웃 정보 가져오기
        first_section = doc.sections[0]
        page_width = first_section.page_width.cm
        page_height = first_section.page_height.cm
        margin_top = first_section.top_margin.cm if first_section.top_margin else 2.54
        margin_bottom = first_section.bottom_margin.cm if first_section.bottom_margin else 2.54
        margin_left = first_section.left_margin.cm if first_section.left_margin else 2.54
        margin_right = first_section.right_margin.cm if first_section.right_margin else 2.54
        orientation = 'landscape' if first_section.orientation == 1 else 'portrait'
        
        # HTML 시작
        html_output = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{html.escape(str(metadata['title']))}</title>
    <meta name="author" content="{html.escape(str(metadata['author']))}">
    <meta name="last-modified" content="{metadata['modified']}">
    <style>
    @charset "UTF-8";
    
    /* 기본 스타일 */
    :root {{
        --main-font: 'Calibri', 'Malgun Gothic', sans-serif;
        --main-color: #333;
        --page-width: {page_width}cm;
        --page-height: {page_height}cm;
        --margin-top: {margin_top}cm;
        --margin-bottom: {margin_bottom}cm;
        --margin-left: {margin_left}cm;
        --margin-right: {margin_right}cm;
        --content-width: calc(var(--page-width) - var(--margin-left) - var(--margin-right));
        --page-orientation: {orientation};
    }}
    
    body {{ 
        font-family: var(--main-font);
        font-size: 11pt;
        line-height: 1.15;
        color: var(--main-color);
    }}
    
    /* 문서 컨테이너 */
    .word-document {{ 
        width: var(--page-width);
        min-height: var(--page-height);
        padding: var(--margin-top) var(--margin-right) var(--margin-bottom) var(--margin-left);
        margin: 1cm auto;
        background: white;
        box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
        position: relative;
        box-sizing: border-box;
    }}
    
    /* 단락 스타일 */
    p {{
        margin: 0;
        padding: 0;
        word-wrap: break-word;
        margin-bottom: 8pt;
    }}
    
    /* 제목 스타일 */
    h1 {{ font-size: 16pt; margin: 24pt 0 12pt 0; }}
    h2 {{ font-size: 14pt; margin: 20pt 0 10pt 0; }}
    h3 {{ font-size: 12pt; margin: 16pt 0 8pt 0; }}
    h4 {{ font-size: 11pt; margin: 14pt 0 7pt 0; }}
    
    /* 표 스타일 */
    table {{ 
        margin: 12pt 0;
        border-collapse: collapse;
        table-layout: fixed;
    }}
    
    td {{ 
        padding: 6pt;
        border: 1px solid #ddd;
        word-wrap: normal;
        white-space: nowrap;
        overflow: visible;
        vertical-align: top;
        box-sizing: border-box;
    }}
    
    /* 이미지 */
    img {{ 
        max-width: 100%;
        height: auto;
        margin: 8pt 0;
        display: block;
    }}
    
    /* 인쇄용 스타일 */
    @media print {{
        @page {{
            size: {orientation};
            margin: 0;
        }}
        
        body {{
            margin: 0;
            padding: 0;
            background: none;
        }}
        
        .word-document {{
            width: 100%;
            min-height: 100%;
            padding: var(--margin-top) var(--margin-right) var(--margin-bottom) var(--margin-left);
            margin: 0;
            box-shadow: none;
            box-sizing: border-box;
            print-color-adjust: exact;
            -webkit-print-color-adjust: exact;
        }}
        
        /* 페이지 나누기 */
        .page-break {{
            page-break-before: always;
        }}
    }}
    </style>
</head>
<body>
{header_footer_html}
<div class="word-document" role="document">
"""
        
        # 문서 내용 처리
        bookmark_refs = {}
        
        # 본문 처리
        for element in doc.element.body:
            if isinstance(element, CT_P):
                # 단락 처리
                paragraph = docx.text.paragraph.Paragraph(element, doc._body)
                html_output += process_paragraph(paragraph, doc)
            elif isinstance(element, CT_Tbl):
                # 표 처리
                table = Table(element, doc._body)
                html_output += process_table(table)
            elif element.tag.endswith('bookmarkStart'):
                # 책갈피 시작
                bookmark_id = element.get('w:id')
                bookmark_name = element.get('w:name', '')
                if bookmark_name:  # 책갈피 이름이 있는 경우에만 처리
                    bookmark_refs[bookmark_id] = bookmark_name
                    html_output += f'<a id="{html.escape(bookmark_name)}"></a>'
            elif element.tag.endswith('comment'):
                # 주석
                html_output += process_comment(element)
            elif element.tag.endswith('ins') or element.tag.endswith('del'):
                # 변경 내역
                html_output += process_revision(element)
            elif element.tag.endswith('drawing'):
                # 이미지
                image_counter += 1
                html_output += process_image(doc, element, image_counter, images_dir)
            elif element.tag.endswith('oMath'):
                # 수식
                html_output += process_equation(element)
        
        # HTML 종료
        html_output += """
</div>
<script>
// 접근성 및 상호작용 개선을 위한 JavaScript
document.addEventListener('DOMContentLoaded', function() {
    // 주석 토글
    const comments = document.querySelectorAll('.comment-marker');
    comments.forEach(comment => {
        comment.addEventListener('click', function() {
            const text = this.nextElementSibling;
            text.style.display = text.style.display === 'block' ? 'none' : 'block';
        });
    });
    
    // 이미지 alt 텍스트 보완
    const images = document.querySelectorAll('img:not([alt])');
    images.forEach((img, index) => {
        img.alt = `문서 이미지 ${index + 1}`;
    });
});
</script>
</body>
</html>
"""
        
        return html_output
        
    except Exception as e:
        logger.error(f"문서 변환 중 오류 발생: {str(e)}")
        raise

def docx_to_html_main(file_paths):
    try:
        if not file_paths or len(file_paths) < 1:
            raise ValueError("최소 1개 이상의 파일 경로가 필요합니다.")
        
        # 첫 번째 파일 경로 사용
        file_path = file_paths[0]
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")
        
        # HTML 변환
        html_content = read_docx_as_html_structure(file_path)
        os.path.splitext(os.path.basename(file_paths[0]))[0] + "hi.html"
        
        return html_content
        
    except Exception as e:
        logger.error(f"오류 발생: {str(e)}")
        raise