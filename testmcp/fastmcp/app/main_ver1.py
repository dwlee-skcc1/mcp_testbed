from mcp.server.fastmcp import FastMCP

import os
import datetime
from docx import Document
import openpyxl  
import PyPDF2
from typing import Literal
from datetime import datetime, timedelta

mcp = FastMCP("test")

@mcp.tool()
def add(a: int, b: int) -> int:
    """Add two numbers"""
    return a + b

@mcp.tool()
def multiply(a: int, b: int) -> int:
    """Multiply two numbers"""
    return a * b


@mcp.tool()
def search_latest_documents(
    keywords: list = None,
    doc_type: Literal['docx', 'xlsx', 'pdf'] = None,
    author: str = None,
    date_range_years: int = None,
    docs_num: int = 1,  
    sort_order: str = "desc", 
    path: str = None,
    search_in_content: bool = False  # 파일 내용 검색 여부
):
    """
    특정 키워드와 조건을 기반으로 최신 문서를 검색하는 함수

    :param keywords: 검색할 키워드들
    :param doc_type: 문서 형식 (예: 'docx', 'xlsx', 'pdf')
    :param modified_by: 문서의 마지막 수정자
    :param date_range_years: 최근 몇 년 간의 문서만 검색
    :param docs_num: 반환할 문서의 개수
    :param sort_order: 정렬 기준 ('asc', 'desc')
    :param path: 문서 검색 경로
    """
    
    path = "C:\\Users\\Administrator\\Desktop\\ye\\LKM\\Tools\\mcp_testbed\\testmcp\\mcpclient\\app\\data"
    
    # 1. 파일 목록 가져오기
    if not path or not os.path.exists(path):
        return {"status": "error", "message": "유효한 경로가 없습니다."}
    
    # 2. 파일 타입 설정
    #######RFQ 특화########
    if 'RFQ' in keywords: 
        doc_type = 'docx'
    elif 'TCS' in keywords:
        doc_type = 'xlsx'
    else:
        return {"status": "error", "message": "파일 종류를 특정하지 못했습니다."}

    
    # 3. 디렉토리 내 파일 필터링
    documents = []
    for root, _, files in os.walk(path):
        for file_name in files:
            # 3-1. doc_type으로 끝나는 문서가 아니면 스킵
            if doc_type and not file_name.lower().endswith(f".{doc_type.lower()}"):
                continue
            
            # 3-2. 키워드가 파일명에 없으면 스킵
            # if keywords and not any(keyword.lower() in file_name.lower() for keyword in keywords):
            #     continue  
            # 3-2. 모든 키워드가 파일명에 포함
            if keywords and not all(keyword.lower() in file_name.lower() for keyword in keywords):
                continue
                
                        
            
            # 3-3. 파일 형식에 따라 메타데이터 추출
            file_path = os.path.join(root, file_name)
            try:
                modified_time = None
                modified_by = "Unknown"
                
                if file_name.lower().endswith('.docx'):
                    doc = Document(file_path)
                    if doc.core_properties.modified:
                        modified_time = doc.core_properties.modified
                    if doc.core_properties.last_modified_by:
                        modified_by = doc.core_properties.last_modified_by
                
                elif file_name.lower().endswith('.xlsx'):
                    wb = openpyxl.load_workbook(file_path, read_only=True)
                    if wb.properties.modified:
                        modified_time = wb.properties.modified
                    if wb.properties.lastModifiedBy:
                        modified_by = wb.properties.lastModifiedBy
                
                elif file_name.lower().endswith('.pdf'):
                    with open(file_path, 'rb') as f:
                        pdf = PyPDF2.PdfReader(f)
                        info = pdf.metadata
                        if info and '/ModDate' in info:
                            # PDF 수정 날짜 포맷 처리 (D:20230101120000+00'00' 형식)
                            mod_date = info['/ModDate'][2:16]  # D:20230101120000 부분 추출
                            try:
                                modified_time = datetime.strptime(mod_date, '%Y%m%d%H%M%S')
                            except:
                                pass
                        if info and '/Author' in info:
                            modified_by = info['/Author']
                
                # 메타데이터에서 수정 시간을 가져오지 못했을 경우 파일 시스템 정보 사용
                if not modified_time:
                    file_stats = os.stat(file_path)
                    modified_time = datetime.fromtimestamp(file_stats.st_mtime)
            
                # 3-4.수정자 필터링
                if author and (not modified_by or author.lower() not in modified_by.lower()):
                    continue
                
                # 3-5. 날짜 필터링
                if date_range_years:
                    cutoff_date = datetime.now() - timedelta(days=365 * date_range_years)
                    if modified_time < cutoff_date:
                        continue
                
                # 3-6. 문서 리스트 추가
                documents.append({
                    "file_name": file_name,
                    "path": file_path,
                    "modified_time": modified_time,
                    "modified_by": modified_by
                })

                print("\n문서 추가")
                print(keywords)
                print(file_name)
                print(modified_time)
                print(modified_by)
                
            except Exception as e:
                print(f"파일 {file_name} 처리 중 오류 발생: {str(e)}")

    # 4. 수정 날짜 sorting
    if sort_order.lower() == 'desc': #최신순
        documents.sort(key=lambda x: x["modified_time"], reverse=True)
    else:
        documents.sort(key=lambda x: x["modified_time"])
    
    # 5. docs_num 개수만큼 선택
    result = [{
        "file_name": doc["file_name"],
        "path": doc["path"],
        "modified_time": doc["modified_time"].isoformat() if doc["modified_time"] else None,
        "modified_by": doc["modified_by"],
        "file_type": doc["file_name"].split('.')[-1] if '.' in doc["file_name"] else None
    } for doc in documents[:docs_num]]
    print(f"\n{docs_num}개 찾기")
    print(f"{len(result)}개 문서: {result}")
    
    return result


def check_keywords_in_docx(doc, keywords, content_pages=2):
    """DOCX 파일 내용에서 키워드 검색"""
    paragraphs_to_check = min(len(doc.paragraphs), content_pages * 10)  # 대략 페이지당 10개 단락으로 가정
    
    for i in range(paragraphs_to_check):
        text = doc.paragraphs[i].text.lower()
        if all(keyword.lower() in text for keyword in keywords):
            return True
    return False

def check_keywords_in_xlsx(workbook, keywords, content_pages=2):
    """XLSX 파일 내용에서 키워드 검색"""
    sheets_to_check = min(len(workbook.sheetnames), content_pages)
    
    for i in range(sheets_to_check):
        sheet = workbook[workbook.sheetnames[i]]
        # 시트의 처음 20x20 셀 영역 검색 (일반적인 헤더 영역)
        for row in range(1, 21):
            for col in range(1, 21):
                cell = sheet.cell(row=row, column=col)
                if cell.value and isinstance(cell.value, str):
                    text = cell.value.lower()
                    if all(keyword.lower() in text for keyword in keywords):
                        return True
    return False

def check_keywords_in_pdf(pdf, keywords, content_pages=2):
    """PDF 파일 내용에서 키워드 검색"""
    pages_to_check = min(len(pdf.pages), content_pages)
    
    for i in range(pages_to_check):
        page = pdf.pages[i]
        text = page.extract_text().lower()
        if all(keyword.lower() in text for keyword in keywords):
            return True
    return False




if __name__ == "__main__":
    mcp.run(transport='sse')